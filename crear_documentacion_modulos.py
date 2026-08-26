from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"C:\Users\isabe\Documents\ChatGPT\Llantas\Documentacion_funcional_GLLD_por_modulos.docx")
NAVY = "123F5C"; BLUE = "0C4B78"; GREEN = "8BC53F"; PALE = "EAF2F5"; GRAY = "607783"; LIGHT = "F4F7F8"; WHITE="FFFFFF"; AMBER="FFF2CC"

modules = [
 ("Inicio / Tablero", "Parcialmente funcional", "Presenta el resumen operativo general del sistema y accesos rápidos a los procesos principales.",
  ["Indicadores generales de llantas y actividades.", "Accesos a módulos operativos.", "Resumen visual para orientar al usuario."],
  ["Los indicadores deben alimentarse desde SQL Server.", "Definir indicadores distintos por rol y centro."], "Dashboard, inspecciones, movimientos, alertas y actividades."),
 ("Inicio de sesión y perfiles", "Prototipo funcional", "Controla el ingreso básico y diferencia los perfiles Administrador, Supervisor y Técnico.",
  ["Ingreso a la aplicación.", "Protección de rutas.", "Restricción básica por rol."],
  ["Conectar con usuarios reales de la empresa.", "Implementar permisos configurables y almacenamiento seguro de contraseñas.", "Incluir el permiso INSPECCION_AUTORIZAR_INCONSISTENCIA_LLANTA."], "Usuarios, roles, permisos, auditoría y centros."),
 ("Mis actividades", "Prototipo funcional", "Muestra al técnico las inspecciones, montajes y revisiones que tiene programadas.",
  ["Listado por fecha, hora, prioridad y estado.", "Inicio de una actividad desde la tarjeta.", "Navegación al proceso correspondiente."],
  ["Guardar programación y ejecución en base de datos.", "Filtrar por técnico, centro, fecha y estado."], "Actividades programadas, vehículos, centros y usuarios."),
 ("Llantas", "Funcional con datos locales/API", "Permite registrar y consultar las llantas que forman el inventario oficial.",
  ["Código y serial.", "Marca, referencia, dimensión y tipo.", "Estado, centro y ubicación.", "Profundidad inicial, costo, fechas y observaciones."],
  ["Conectar completamente con SQL Server.", "Validar códigos y seriales únicos.", "Incluir documentos, fotografías e historial completo."], "Marcas, referencias, dimensiones, tipos, estados, centros e historial."),
 ("Vehículos y posiciones", "Pantalla demostrativa", "Administra la flota y la configuración variable de ejes y posiciones de cada vehículo.",
  ["Número interno y placa.", "Tipo, marca y modelo.", "Centro de trabajo.", "Ejes, posiciones y llanta actualmente instalada."],
  ["Crear editor gráfico de ejes y posiciones.", "No asumir una cantidad fija de llantas.", "Validar una asignación activa por llanta y posición."], "Vehículos, centros, ejes, posiciones, asignaciones y movimientos."),
 ("Inventario", "Pantalla demostrativa", "Consulta la disponibilidad y ubicación de las llantas por centro y estado.",
  ["Búsqueda por llanta.", "Filtros por centro, marca, dimensión y estado.", "Identificación de llantas montadas y disponibles."],
  ["Conectar existencias reales.", "Restringir la visualización según centros autorizados.", "Incluir ubicación física de bodega."], "Llantas, centros, estados, movimientos y ubicaciones."),
 ("Inspecciones", "Prototipo funcional avanzado", "Registra una inspección completa del vehículo conservando la relación vehículo, eje, posición y llanta.",
  ["Selección por interno y placa.", "Centro asociado mostrado por nombre; R1/R2/R3/R4 queda interno.", "Kilometraje vacío al iniciar.", "Esquema adaptable de ejes y todas las posiciones.", "Profundidades Exterior, Centro e Interior en desplegables.", "Condición/novedad, causa, observaciones y recomendación.", "Reporte de llanta diferente y creación de inconsistencia temporal."],
  ["Consumir vehículos, posiciones y catálogos desde SQL Server.", "Guardar borradores y finalizar inspecciones.", "Subir evidencias reales.", "Aplicar autorización de inconsistencias sin permitir autoaprobación.", "Parametrizar reglas de reencauche y kilometraje."], "Inspecciones, detalles, vehículos, posiciones, llantas, condiciones, causas, recomendaciones, inconsistencias, temporales y evidencias."),
 ("Alertas y autorizaciones", "Pantalla demostrativa / backend preparado", "Centraliza hallazgos e inconsistencias que requieren revisión o autorización.",
  ["Alerta por llanta esperada diferente de la encontrada.", "Consulta de solicitudes pendientes.", "Autorización o rechazo con observación.", "Cierre o actualización de la alerta."],
  ["Crear bandeja específica para Visor Técnico.", "Aplicar permisos por usuario y centro.", "Registrar auditoría de todas las decisiones."], "Inconsistencias, llantas temporales, evidencias, usuarios, permisos y auditoría."),
 ("Programación", "Pantalla demostrativa", "Programa inspecciones, rotaciones, cambios y mantenimientos.",
  ["Tipo de actividad.", "Vehículo, fecha, técnico y prioridad.", "Seguimiento del estado."],
  ["Calendario operativo.", "Reglas para actividades vencidas.", "Notificaciones y reasignaciones."], "Actividades, vehículos, técnicos, centros y alertas."),
 ("Montajes y desmontajes", "Prototipo funcional", "Controla la instalación o retiro de una llanta y genera trazabilidad del movimiento.",
  ["Selección de vehículo y posición.", "Selección de llanta para montaje.", "Destino de la llanta desmontada.", "Motivo, técnico, kilometraje y observaciones."],
  ["Validar disponibilidad y compatibilidad de dimensión.", "Actualizar asignaciones y estados en una transacción.", "Solicitar autorización para movimientos especiales."], "Llantas, posiciones, asignaciones, movimientos, centros e historial."),
 ("Movimientos", "Prototipo funcional", "Registra traslados entre posiciones, inventario, reparación, reencauche, disposición final u otro centro.",
  ["Origen y destino.", "Tratamiento de posición ocupada.", "Motivo y resumen previo.", "Movimiento con detalle e historial."],
  ["Conectar endpoints y SQL Server.", "Evitar actualizaciones directas de posición.", "Aplicar transacciones y concurrencia."], "Movimientos, detalles, asignaciones, llantas, posiciones y centros."),
 ("Reparaciones", "Pantalla demostrativa", "Gestiona diagnóstico, proveedor, costos, evidencias y resultado de reparación.",
  ["Envío a reparación.", "Seguimiento del servicio.", "Recepción y resultado técnico."],
  ["Definir catálogo de daños y talleres.", "Registrar costos y garantía.", "Integrar movimientos de salida y entrada."], "Llantas, talleres, reparaciones, evidencias, costos y movimientos."),
 ("Reencauches", "Pantalla demostrativa / reglas preparadas", "Controla la evaluación, envío, recepción y nueva vida útil de una llanta reencauchada.",
  ["Recomendación desde inspección sin movimiento automático.", "Evaluación de carcasa.", "Registro de banda y número de reencauche."],
  ["Definir parámetros por dimensión y vigencia.", "Registrar proveedor, costos, fechas, garantía y resultado.", "Actualizar profundidad y ciclo de vida al recibir."], "Parámetros, dimensiones, inspecciones, movimientos, proveedores y llantas."),
 ("Disposición final", "Prototipo funcional", "Controla la salida definitiva de llantas y genera el acta en PDF.",
  ["R2, R3 y R4 envían a R1.", "R1 recibe conservando la planta de origen.", "Solo R1 entrega a la empresa de disposición.", "Selección de llantas y descarga del acta PDF."],
  ["Guardar actas y firmas.", "Agregar consecutivo y empresa receptora.", "Bloquear entrega de llantas que no estén recibidas en R1."], "Llantas, centros, movimientos, actas, detalles, evidencias y auditoría."),
 ("Historial de llantas", "Pantalla demostrativa", "Presenta cronológicamente todo lo ocurrido a una llanta.",
  ["Inspecciones.", "Montajes, desmontajes y traslados.", "Reparaciones, reencauches y disposición final."],
  ["Construir una línea de tiempo consolidada desde SQL Server.", "Permitir búsqueda por código, serial, vehículo y fecha."], "Todas las tablas transaccionales y de auditoría."),
 ("Carga masiva", "Pantalla demostrativa", "Permite importar información desde Excel con validación previa.",
  ["Selección del tipo de carga.", "Previsualización.", "Conteo de registros correctos y rechazados.", "Reporte de errores."],
  ["Definir plantillas oficiales.", "Evitar duplicados.", "Procesar en lotes y registrar quién realizó la carga."], "Catálogos, llantas, vehículos, centros y auditoría de cargas."),
 ("Analítica", "Pantalla demostrativa", "Presenta indicadores de rendimiento, costos, desgaste y comportamiento de marcas.",
  ["Indicadores generales.", "Tendencias de desgaste.", "Apoyo al análisis de costo por kilómetro y marcas."],
  ["Definir fórmulas oficiales de CPK.", "Agregar filtros y exportación.", "Validar calidad de datos antes de publicar indicadores."], "Inspecciones, movimientos, costos, reparaciones, reencauches, vehículos y llantas."),
 ("Administración y catálogos", "Prototipo funcional", "Administra parámetros usados en los demás módulos.",
  ["Marcas, referencias, dimensiones, tipos y estados.", "Centros, talleres y técnicos.", "Motivos, tolerancias y reglas de reencauche."],
  ["Conectar todos los catálogos con SQL Server.", "Controlar vigencias y evitar borrados con información histórica.", "Administrar los 151 centros y su relevancia interna."], "Todos los catálogos, parámetros, usuarios y auditoría."),
 ("Auditoría", "Pantalla demostrativa / base preparada", "Permite conocer quién cambió qué, cuándo y con qué resultado.",
  ["Usuario, fecha y operación.", "Entidad y registro afectado.", "Valores o resultado de la acción."],
  ["Registrar inicio de sesión, autorizaciones y cambios críticos.", "Permitir consulta y exportación sin alterar los registros."], "Usuarios y todas las tablas auditables."),
]

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def margins(cell, top=100, start=140, bottom=100, end=140):
    tc=cell._tc.get_or_add_tcPr(); m=tc.first_child_found_in('w:tcMar')
    if m is None: m=OxmlElement('w:tcMar'); tc.append(m)
    for tag,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        e=OxmlElement('w:'+tag); e.set(qn('w:w'),str(val)); e.set(qn('w:type'),'dxa'); m.append(e)
def add_cell_text(cell,text,bold=False,color=NAVY,size=9):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(text); r.bold=bold;r.font.name='Calibri';r.font.size=Pt(size);r.font.color.rgb=RGBColor.from_string(color);cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER;margins(cell)
def bullet(doc,text):
    p=doc.add_paragraph(style='List Bullet');p.paragraph_format.space_after=Pt(3);p.paragraph_format.line_spacing=1.15;p.add_run(text);return p
def response_box(doc):
    t=doc.add_table(rows=3,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
    labels=['MEJORA SOLICITADA:','PRIORIDAD / RESPONSABLE:','CRITERIO PARA CONSIDERARLA TERMINADA:']
    for i,c in enumerate(t.column_cells(0)):
        c.width=Inches(6.5);shade(c,'F7F9FA');add_cell_text(c,labels[i]+'\n\n',True,GRAY,9)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

doc=Document(); sec=doc.sections[0];sec.page_width=Inches(8.5);sec.page_height=Inches(11);sec.top_margin=sec.bottom_margin=Inches(.8);sec.left_margin=sec.right_margin=Inches(.85);sec.header_distance=Inches(.35);sec.footer_distance=Inches(.35)
styles=doc.styles
normal=styles['Normal'];normal.font.name='Calibri';normal.font.size=Pt(10.5);normal.font.color.rgb=RGBColor.from_string('263B45');normal.paragraph_format.space_after=Pt(6);normal.paragraph_format.line_spacing=1.15
for name,size,before,after,color in [('Title',28,0,8,NAVY),('Heading 1',18,14,7,BLUE),('Heading 2',14,12,5,NAVY),('Heading 3',11,8,4,GREEN)]:
    s=styles[name];s.font.name='Calibri';s.font.size=Pt(size);s.font.bold=True;s.font.color.rgb=RGBColor.from_string(color);s.paragraph_format.space_before=Pt(before);s.paragraph_format.space_after=Pt(after);s.paragraph_format.keep_with_next=True
header=sec.header.paragraphs[0];header.text='GLLD  |  EDINSA';header.style=styles['Normal'];header.runs[0].font.bold=True;header.runs[0].font.color.rgb=RGBColor.from_string(GRAY);header.runs[0].font.size=Pt(8)
footer=sec.footer.paragraphs[0];footer.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=footer.add_run('Documento de revisión funcional por módulos');r.font.name='Calibri';r.font.size=Pt(8);r.font.color.rgb=RGBColor.from_string(GRAY)

p=doc.add_paragraph();p.paragraph_format.space_before=Pt(70);p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('GLLD');r.bold=True;r.font.name='Calibri';r.font.size=Pt(40);r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('Sistema de Gestión Logística de Llantas');r.font.size=Pt(17);r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(14);r=p.add_run('Documentación funcional y formato para solicitar mejoras por módulo');r.bold=True;r.font.size=Pt(13);r.font.color.rgb=RGBColor.from_string(GREEN)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(80);r=p.add_run('EDINSA · Transportadora de Postobón');r.bold=True;r.font.size=Pt(12);r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('Versión de trabajo · 12 de agosto de 2026');r.font.size=Pt(10);r.font.color.rgb=RGBColor.from_string(GRAY)
doc.add_page_break()

doc.add_heading('Cómo utilizar este documento',level=1)
doc.add_paragraph('Este archivo resume el alcance actual de cada módulo. En los recuadros de mejora puedes escribir exactamente lo que deseas cambiar y devolver el documento para convertir esas observaciones en un plan de implementación.')
doc.add_heading('Estados usados',level=2)
t=doc.add_table(rows=1,cols=2);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
for c,txt in zip(t.rows[0].cells,['Estado','Significado']):shade(c,NAVY);add_cell_text(c,txt,True,WHITE,9)
for status,meaning in [('Funcional con datos locales/API','Cuenta con operaciones implementadas, pero todavía debe verificarse la conexión empresarial.'),('Prototipo funcional','Permite demostrar el flujo y las interacciones principales.'),('Prototipo funcional avanzado','Incluye una experiencia detallada y reglas importantes aún por conectar con datos reales.'),('Pantalla demostrativa','Muestra el concepto del módulo; requiere completar backend y persistencia.'),('Backend preparado','Existen entidades o servicios base, pero falta completar la experiencia integral.')]:
    cells=t.add_row().cells;add_cell_text(cells[0],status,True,NAVY,8.5);add_cell_text(cells[1],meaning,False,'263B45',8.5)
doc.add_heading('Reglas generales que aplican a todos los módulos',level=2)
for x in ['El centro de trabajo se muestra por nombre; R1, R2, R3 y R4 se conservan como clasificación interna.','No se debe perder el historial de una llanta, posición, vehículo o movimiento.','Los catálogos deben ser configurables y no quedar fijos en el código.','Las acciones críticas deben registrar usuario, fecha, observación y resultado.','La interfaz debe funcionar en computador, tableta y celular.']:bullet(doc,x)

for idx,(name,status,purpose,current,pending,relations) in enumerate(modules,1):
    doc.add_page_break()
    doc.add_heading(f'{idx}. {name}',level=1)
    table=doc.add_table(rows=1,cols=2);table.alignment=WD_TABLE_ALIGNMENT.CENTER;table.autofit=False
    table.columns[0].width=Inches(1.6);table.columns[1].width=Inches(4.9)
    c=table.rows[0].cells;shade(c[0],NAVY);shade(c[1],PALE);add_cell_text(c[0],'ESTADO ACTUAL',True,WHITE,9);add_cell_text(c[1],status,True,NAVY,10)
    doc.add_heading('Objetivo',level=2);doc.add_paragraph(purpose)
    doc.add_heading('Qué incluye actualmente',level=2)
    for item in current:bullet(doc,item)
    doc.add_heading('Pendientes conocidos',level=2)
    for item in pending:bullet(doc,item)
    doc.add_heading('Información relacionada',level=2);doc.add_paragraph(relations)
    doc.add_heading('Espacio para solicitar mejoras',level=2);response_box(doc)

doc.add_page_break();doc.add_heading('Formato para una mejora transversal',level=1)
doc.add_paragraph('Utiliza esta sección cuando una solicitud afecte más de un módulo.')
for label in ['MÓDULOS AFECTADOS:','DESCRIPCIÓN DE LA NECESIDAD:','USUARIO O ROL QUE LA UTILIZA:','FLUJO ESPERADO PASO A PASO:','DATOS OBLIGATORIOS:','VALIDACIONES Y AUTORIZACIONES:','RESULTADO ESPERADO:','EJEMPLO, ARCHIVO O IMAGEN DE REFERENCIA:']:
    t=doc.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;c=t.cell(0,0);shade(c,LIGHT);add_cell_text(c,label+'\n\n',True,GRAY,9);doc.add_paragraph().paragraph_format.space_after=Pt(0)

doc.core_properties.title='Documentación funcional GLLD por módulos';doc.core_properties.subject='Formato de revisión y solicitud de mejoras';doc.core_properties.author='GLLD - EDINSA';doc.save(OUT)
print(OUT)
